from langchain_community.document_loaders import (
    PyMuPDFLoader,
    WebBaseLoader,
    YoutubeLoader
)
from langchain.docstore.document import Document
import tempfile
import os

def load_documents(files=None, url=None):
    docs = []

    if files:
        for file in files:
            suffix = file.name.split(".")[-1].lower()
            print(f"Processing file: {file.name} ({suffix})")
            try:
                with tempfile.NamedTemporaryFile(delete=False, suffix=f".{suffix}") as tmp:
                    tmp.write(file.read())
                    tmp_path = tmp.name

                # Safer DOCX
                if suffix == "docx":
                    from docx import Document as DocxDocument
                    text = "\n".join([p.text for p in DocxDocument(tmp_path).paragraphs])
                    docs.append(Document(page_content=text, metadata={"source": file.name}))

                # Safer PDF (fallback to PyMuPDF)
                elif suffix == "pdf":
                    try:
                        loader = PyMuPDFLoader(tmp_path)
                        file_docs = loader.load()
                        docs.extend(file_docs)
                        print(f"Loaded PDF with {len(file_docs)} pages")
                    except Exception as e:
                        print(f"Failed to load PDF: {e}")


                # Safer Image (fallback to Tesseract OCR)
                elif suffix in ["jpg", "jpeg", "png"]:
                    try:
                        from PIL import Image
                        import pytesseract
                        pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"
                        image = Image.open(tmp_path)
                        text = pytesseract.image_to_string(image)
                        docs.append(Document(page_content=text, metadata={"source": file.name}))
                        print(f"Image processed with Tesseract OCR: {file.name}")
                    except Exception as e:
                        print(f"Failed to OCR image {file.name}: {e}")

                else:
                    print(f"Unsupported file type: {suffix}")
                    continue

            except Exception as e:
                print(f"Error processing file {file.name}: {e}")

    # Handle URL input
    if url:
        print(f"Loading from URL: {url}")
        try:
            if "youtube.com" in url or "youtu.be" in url:
                try:
                    from pytube import YouTube

                    yt = YouTube(url)
                    caption_text = ""

                    # Prefer English caption
                    caption = yt.captions.get("en") or yt.captions.get("a.en") or yt.captions.get_by_language_code("en")
                    if caption:
                        caption_text = caption.generate_srt_captions()
                    else:
                        print("No captions found, using video description")
                        caption_text = yt.description

                    docs.append(Document(page_content=caption_text, metadata={"source": url}))
                    print(f"Loaded YouTube transcript/description from: {url}")

                except Exception as e:
                    print(f"Failed to load YouTube video: {e}")
                    from yt_dlp import YoutubeDL
                    from bs4 import BeautifulSoup

                    with YoutubeDL({'quiet': True, 'skip_download': True}) as ydl:
                        info = ydl.extract_info(url, download=False)
                        subtitles = info.get("subtitles", {})
                        transcript = ""

                        # try grabbing transcript from subs
                        for lang in ["en", "en-US", "auto"]:
                            if lang in subtitles:
                                for track in subtitles[lang]:
                                    if track["ext"] == "vtt":
                                        import requests
                                        vtt_url = track["url"]
                                        response = requests.get(vtt_url)
                                        transcript = response.text
                                        break
                            if transcript:
                                break

                        if not transcript:
                            transcript = info.get("description", "")

                        docs.append(Document(page_content=transcript, metadata={"source": url}))
            else:
                try:
                    import requests
                    from bs4 import BeautifulSoup

                    headers = {"User-Agent": os.getenv("USER_AGENT", "docubot")}
                    response = requests.get(url, headers=headers)
                    soup = BeautifulSoup(response.text, "html.parser")

                    # Try to extract <main> if available, else fallback to full body text
                    main_content = soup.find("main")
                    text = main_content.get_text(separator="\n") if main_content else soup.get_text(separator="\n")

                    text = text.strip()
                    if text:
                        docs.append(Document(page_content=text, metadata={"source": url}))
                        print(f"Loaded {len(text.split())} words from webpage.")
                    else:
                        print("Webpage loaded but returned no readable content.")
                except Exception as e:
                    print(f"Error parsing webpage manually: {e}")


        except Exception as e:
            print(f"Error loading URL: {e}")

    return docs
